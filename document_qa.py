# """
# Document QA Module

# This module provides document question-answering functionality.
# It handles document storage, indexing, and retrieval.
# """

# import os
# import logging
# import json
# import shutil
# from typing import Dict, Any, List, Optional, Tuple
# import uuid
# from datetime import datetime

# # Configure logging
# logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
# logger = logging.getLogger(__name__)

# # Try to import Elasticsearch
# try:
#     from elasticsearch import Elasticsearch
#     from elasticsearch.exceptions import ConnectionError as ESConnectionError
#     ELASTICSEARCH_AVAILABLE = True
# except ImportError:
#     logger.warning("Elasticsearch Python client not available. Using local file-based search.")
#     ELASTICSEARCH_AVAILABLE = False

# # Global variables
# UPLOAD_DIR = None
# ES_CLIENT = None
# ES_INDEX = "document_qa"

# def init_document_qa(upload_dir: str) -> None:
#     """Initialize the Document QA module."""
#     global UPLOAD_DIR, ES_CLIENT
    
#     UPLOAD_DIR = upload_dir
#     logger.info(f"Document QA initialized with upload directory: {UPLOAD_DIR}")
    
#     # Create user documents directory if it doesn't exist
#     os.makedirs(os.path.join(UPLOAD_DIR, "user_documents"), exist_ok=True)
    
#     # Initialize Elasticsearch if available
#     if ELASTICSEARCH_AVAILABLE:
#         try:
#             ES_CLIENT = Elasticsearch(["http://localhost:9200"])
#             if not ES_CLIENT.ping():
#                 logger.warning("Elasticsearch server not responding. Using local file-based search.")
#                 ES_CLIENT = None
#             else:
#                 # Create index if it doesn't exist
#                 if not ES_CLIENT.indices.exists(index=ES_INDEX):
#                     ES_CLIENT.indices.create(
#                         index=ES_INDEX,
#                         body={
#                             "mappings": {
#                                 "properties": {
#                                     "content": {"type": "text"},
#                                     "file_name": {"type": "keyword"},
#                                     "file_path": {"type": "keyword"},
#                                     "file_size": {"type": "long"},
#                                     "upload_date": {"type": "date"},
#                                     "user_id": {"type": "keyword"}
#                                 }
#                             }
#                         }
#                     )
#                 logger.info("Elasticsearch connected and index checked.")
#         except Exception as e:
#             logger.warning(f"Elasticsearch server not responding. Using local file-based search. Error: {type(e).__name__}")
#             ES_CLIENT = None
#     else:
#         logger.warning("Elasticsearch not available. Using local file-based search.")

# def extract_text_from_document(file_path: str) -> str:
#     """Extract text from a document file."""
#     file_ext = os.path.splitext(file_path)[1].lower()
    
#     if file_ext == '.pdf':
#         return extract_text_from_pdf(file_path)
#     elif file_ext == '.txt':
#         return extract_text_from_txt(file_path)
#     elif file_ext == '.docx':
#         return extract_text_from_docx(file_path)
#     else:
#         logger.warning(f"Unsupported file format: {file_ext}")
#         return ""

# def extract_text_from_pdf(file_path: str) -> str:
#     """Extract text from a PDF file."""
#     try:
#         import PyPDF2
        
#         with open(file_path, 'rb') as file:
#             pdf_reader = PyPDF2.PdfReader(file)
#             text = ""
#             for page_num in range(len(pdf_reader.pages)):
#                 page = pdf_reader.pages[page_num]
#                 text += page.extract_text() + "\n\n"
#             return text
#     except ImportError:
#         logger.error("PyPDF2 not installed. Cannot extract text from PDF.")
#         return ""
#     except Exception as e:
#         logger.error(f"Error extracting text from PDF: {str(e)}")
#         return ""

# def extract_text_from_txt(file_path: str) -> str:
#     """Extract text from a TXT file."""
#     try:
#         with open(file_path, 'r', encoding='utf-8') as file:
#             return file.read()
#     except UnicodeDecodeError:
#         # Try with a different encoding if UTF-8 fails
#         try:
#             with open(file_path, 'r', encoding='latin-1') as file:
#                 return file.read()
#         except Exception as e:
#             logger.error(f"Error reading TXT file with latin-1 encoding: {str(e)}")
#             return ""
#     except Exception as e:
#         logger.error(f"Error reading TXT file: {str(e)}")
#         return ""

# def extract_text_from_docx(file_path: str) -> str:
#     """Extract text from a DOCX file."""
#     try:
#         import docx
        
#         doc = docx.Document(file_path)
#         text = ""
#         for para in doc.paragraphs:
#             text += para.text + "\n"
#         return text
#     except ImportError:
#         logger.error("python-docx not installed. Cannot extract text from DOCX.")
#         return ""
#     except Exception as e:
#         logger.error(f"Error extracting text from DOCX: {str(e)}")
#         return ""

# def add_document(file_path: str, user_id: int, file_size: int = 0) -> Optional[str]:
#     """Add a document to the QA system."""
#     if not os.path.exists(file_path):
#         logger.error(f"File not found: {file_path}")
#         return None
    
#     try:
#         # Generate a unique document ID
#         doc_id = str(uuid.uuid4())
        
#         # Extract text from the document
#         text = extract_text_from_document(file_path)
#         if not text:
#             logger.warning(f"No text extracted from document: {file_path}")
#             return None
        
#         # Get file metadata
#         file_name = os.path.basename(file_path)
#         upload_date = datetime.utcnow().isoformat()
        
#         # Create user document directory if it doesn't exist
#         user_dir = os.path.join(UPLOAD_DIR, "user_documents", str(user_id))
#         os.makedirs(user_dir, exist_ok=True)
        
#         # Copy the document to the user directory
#         dest_path = os.path.join(user_dir, f"{doc_id}_{file_name}")
#         shutil.copy2(file_path, dest_path)
        
#         # Create document metadata
#         doc_metadata = {
#             "id": doc_id,
#             "file_name": file_name,
#             "file_path": dest_path,
#             "file_size": file_size,
#             "upload_date": upload_date,
#             "user_id": str(user_id)
#         }
        
#         # Save document metadata
#         metadata_path = os.path.join(user_dir, f"{doc_id}_metadata.json")
#         with open(metadata_path, 'w') as f:
#             json.dump(doc_metadata, f)
        
#         # Save document content
#         content_path = os.path.join(user_dir, f"{doc_id}_content.txt")
#         with open(content_path, 'w', encoding='utf-8') as f:
#             f.write(text)
        
#         # Index the document in Elasticsearch if available
#         if ES_CLIENT:
#             ES_CLIENT.index(
#                 index=ES_INDEX,
#                 id=doc_id,
#                 body={
#                     "content": text,
#                     "file_name": file_name,
#                     "file_path": dest_path,
#                     "file_size": file_size,
#                     "upload_date": upload_date,
#                     "user_id": str(user_id)
#                 }
#             )
        
#         logger.info(f"Document added: {file_name} (ID: {doc_id})")
#         return doc_id
#     except Exception as e:
#         logger.error(f"Error adding document: {str(e)}")
#         return None

# def query_document(doc_id: str, question: str) -> Dict[str, Any]:
#     """Query a document with a question."""
#     try:
#         # Get document metadata
#         metadata = get_document_metadata(doc_id)
#         if not metadata:
#             logger.error(f"Document not found: {doc_id}")
#             return {"error": "Document not found"}
        
#         user_id = metadata["user_id"]
#         user_dir = os.path.join(UPLOAD_DIR, "user_documents", str(user_id))
#         content_path = os.path.join(user_dir, f"{doc_id}_content.txt")
        
#         if not os.path.exists(content_path):
#             logger.error(f"Document content not found: {doc_id}")
#             return {"error": "Document content not found"}
        
#         # Read document content
#         with open(content_path, 'r', encoding='utf-8') as f:
#             content = f.read()
        
#         # If Elasticsearch is available, use it for search
#         if ES_CLIENT:
#             result = search_with_elasticsearch(doc_id, question)
#             if result:
#                 return result
        
#         # Fallback to simple keyword search
#         result = simple_keyword_search(content, question)
#         if not result or result == "No relevant information found.":
#             return {
#                 "answer": "I couldn't find a specific answer to your question in this document.",
#                 "context": "The document doesn't contain information directly related to your query."
#             }
        
#         return {
#             "answer": result,
#             "context": get_context(content, result)
#         }
#     except Exception as e:
#         logger.error(f"Error querying document: {str(e)}")
#         return {"error": f"Error querying document: {str(e)}"}

# def search_with_elasticsearch(doc_id: str, question: str) -> Optional[Dict[str, Any]]:
#     """Search for an answer using Elasticsearch."""
#     try:
#         # Search for the question in the document
#         response = ES_CLIENT.search(
#             index=ES_INDEX,
#             body={
#                 "query": {
#                     "bool": {
#                         "must": [
#                             {"match": {"content": question}},
#                             {"term": {"_id": doc_id}}
#                         ]
#                     }
#                 },
#                 "highlight": {
#                     "fields": {
#                         "content": {}
#                     }
#                 }
#             }
#         )
        
#         # Extract the best match
#         hits = response["hits"]["hits"]
#         if not hits:
#             return None
        
#         # Get the highlighted fragments
#         highlights = hits[0].get("highlight", {}).get("content", [])
#         if not highlights:
#             return None
        
#         # Join the highlights as the answer
#         answer = " ".join(highlights)
        
#         # Get the document content for context
#         content = hits[0]["_source"]["content"]
        
#         return {
#             "answer": answer,
#             "context": get_context(content, answer)
#         }
#     except Exception as e:
#         logger.error(f"Elasticsearch search error: {str(e)}")
#         return None

# def simple_keyword_search(content: str, question: str) -> str:
#     """Simple keyword-based search for an answer."""
#     # Extract keywords from the question
#     import re
#     from nltk.tokenize import word_tokenize
#     from nltk.corpus import stopwords
    
#     try:
#         # Tokenize and remove stopwords
#         stop_words = set(stopwords.words('english'))
#         question_words = word_tokenize(question.lower())
#         keywords = [word for word in question_words if word.isalnum() and word not in stop_words]
        
#         # Split content into sentences
#         sentences = re.split(r'(?<=[.!?])\s+', content)
        
#         # Score sentences based on keyword matches
#         sentence_scores = []
#         for sentence in sentences:
#             score = 0
#             for keyword in keywords:
#                 if keyword.lower() in sentence.lower():
#                     score += 1
#             if score > 0:
#                 sentence_scores.append((sentence, score))
        
#         # Sort sentences by score
#         sentence_scores.sort(key=lambda x: x[1], reverse=True)
        
#         # Return the top sentences
#         if sentence_scores:
#             top_sentences = [s[0] for s in sentence_scores[:3]]
#             return " ".join(top_sentences)
        
#         return "No relevant information found."
#     except Exception as e:
#         logger.error(f"Simple keyword search error: {str(e)}")
#         return "Error performing search."

# def get_context(content: str, answer: str) -> str:
#     """Get the context surrounding an answer."""
#     # Find the position of the answer in the content
#     pos = content.find(answer)
#     if pos == -1:
#         # If exact match not found, try case-insensitive
#         pos = content.lower().find(answer.lower())
#         if pos == -1:
#             # If still not found, return a portion of the content
#             return content[:500] + "..."
    
#     # Get context before and after the answer
#     start = max(0, pos - 200)
#     end = min(len(content), pos + len(answer) + 200)
    
#     # Adjust to complete sentences
#     while start > 0 and content[start] not in ".!?":
#         start -= 1
#     if start > 0:
#         start += 2  # Skip the period and space
    
#     while end < len(content) and content[end] not in ".!?":
#         end += 1
#     if end < len(content):
#         end += 1  # Include the period
    
#     return content[start:end]

# def get_document_metadata(doc_id: str) -> Optional[Dict[str, Any]]:
#     """Get document metadata."""
#     # Search in all user directories
#     user_docs_dir = os.path.join(UPLOAD_DIR, "user_documents")
#     if not os.path.exists(user_docs_dir):
#         logger.error(f"User documents directory not found: {user_docs_dir}")
#         return None
    
#     for user_id in os.listdir(user_docs_dir):
#         user_dir = os.path.join(user_docs_dir, user_id)
#         if not os.path.isdir(user_dir):
#             continue
        
#         metadata_path = os.path.join(user_dir, f"{doc_id}_metadata.json")
#         if os.path.exists(metadata_path):
#             try:
#                 with open(metadata_path, 'r') as f:
#                     return json.load(f)
#             except Exception as e:
#                 logger.error(f"Error reading document metadata: {str(e)}")
#                 return None
    
#     return None

# def get_user_documents(user_id: int) -> List[Dict[str, Any]]:
#     """Get all documents for a user."""
#     user_dir = os.path.join(UPLOAD_DIR, "user_documents", str(user_id))
#     if not os.path.exists(user_dir):
#         logger.info(f"User directory not found: {user_dir}")
#         return []
    
#     documents = []
#     for filename in os.listdir(user_dir):
#         if filename.endswith("_metadata.json"):
#             try:
#                 with open(os.path.join(user_dir, filename), 'r') as f:
#                     metadata = json.load(f)
#                     # Format the file size
#                     if "file_size" in metadata:
#                         size_bytes = metadata["file_size"]
#                         if size_bytes < 1024:
#                             metadata["file_size_formatted"] = f"{size_bytes} bytes"
#                         elif size_bytes < 1024 * 1024:
#                             metadata["file_size_formatted"] = f"{size_bytes / 1024:.1f} KB"
#                         else:
#                             metadata["file_size_formatted"] = f"{size_bytes / (1024 * 1024):.1f} MB"
#                     else:
#                         metadata["file_size_formatted"] = "Unknown size"
#                     documents.append(metadata)
#             except Exception as e:
#                 logger.error(f"Error reading document metadata: {str(e)}")
    
#     # Sort by upload date (newest first)
#     documents.sort(key=lambda x: x.get("upload_date", ""), reverse=True)
    
#     return documents

# def delete_document(doc_id: str, user_id: int) -> bool:
#     """Delete a document."""
#     try:
#         # Get document metadata
#         metadata = get_document_metadata(doc_id)
#         if not metadata:
#             logger.error(f"Document not found: {doc_id}")
#             return False
        
#         # Check if the document belongs to the user
#         if str(metadata["user_id"]) != str(user_id):
#             logger.error(f"Document {doc_id} does not belong to user {user_id}")
#             return False
        
#         # Delete document files
#         user_dir = os.path.join(UPLOAD_DIR, "user_documents", str(user_id))
#         file_path = metadata.get("file_path")
#         if file_path and os.path.exists(file_path):
#             os.remove(file_path)
        
#         metadata_path = os.path.join(user_dir, f"{doc_id}_metadata.json")
#         if os.path.exists(metadata_path):
#             os.remove(metadata_path)
        
#         content_path = os.path.join(user_dir, f"{doc_id}_content.txt")
#         if os.path.exists(content_path):
#             os.remove(content_path)
        
#         # Delete from Elasticsearch if available
#         if ES_CLIENT:
#             ES_CLIENT.delete(index=ES_INDEX, id=doc_id, ignore=[404])
        
#         logger.info(f"Document deleted: {doc_id}")
#         return True
#     except Exception as e:
#         logger.error(f"Error deleting document: {str(e)}")
#         return False
"""
Document QA Module

This module provides document question-answering functionality.
It handles document storage, indexing, and retrieval with graceful
Elasticsearch connection handling and fallback to local search.
"""

import os
import logging
import json
import shutil
from typing import Dict, Any, List, Optional, Tuple
import uuid
from datetime import datetime

# Configure logging
# logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s') # Configured in app.py
logger = logging.getLogger(__name__)

# Try to import Elasticsearch
try:
    from elasticsearch import Elasticsearch, ConnectionError as ESConnectionError, NotFoundError as ESNotFoundError
    ELASTICSEARCH_AVAILABLE = True
    logger.info("Elasticsearch Python client available.")
except ImportError:
    logger.warning("Elasticsearch Python client not available. Local file-based search will be the only option.")
    ELASTICSEARCH_AVAILABLE = False
    Elasticsearch = None # Define for type hinting if not available
    ESConnectionError = None 
    ESNotFoundError = None

# Global variables
UPLOAD_DIR: Optional[str] = None
ES_CLIENT: Optional[Elasticsearch] = None # Type hint for Elasticsearch client
ES_INDEX: str = "document_qa"
ELASTICSEARCH_URL: str = os.environ.get("ELASTICSEARCH_URL", "http://localhost:9200")

def init_document_qa(upload_dir: str) -> None:
    """Initialize the Document QA module."""
    global UPLOAD_DIR, ES_CLIENT
    
    UPLOAD_DIR = upload_dir
    if not UPLOAD_DIR:
        logger.error("Upload directory not provided for Document QA initialization.")
        return
    logger.info(f"Document QA initializing with upload directory: {UPLOAD_DIR}")
    
    # Create user documents directory if it doesn't exist
    try:
        user_docs_main_path = os.path.join(UPLOAD_DIR, "user_documents")
        os.makedirs(user_docs_main_path, exist_ok=True)
        logger.info(f"Ensured user_documents directory exists at: {user_docs_main_path}")
    except OSError as e:
        logger.error(f"Could not create user_documents directory at {user_docs_main_path}: {e}")
        # Depending on severity, you might want to raise an exception or disable features
        return

    # Initialize Elasticsearch if available and configured
    if ELASTICSEARCH_AVAILABLE and ELASTICSEARCH_URL and ELASTICSEARCH_URL.strip():
        try:
            logger.info(f"Attempting to connect to Elasticsearch at {ELASTICSEARCH_URL}...")
            ES_CLIENT = Elasticsearch(
                [ELASTICSEARCH_URL],
                timeout=10, # Connection timeout
                max_retries=2,
                retry_on_timeout=True
            )
            if not ES_CLIENT.ping():
                logger.warning(f"Elasticsearch ping failed at {ELASTICSEARCH_URL}. Client initialized but may not be responsive. Falling back to local search.")
                ES_CLIENT = None # Set to None if ping fails
            else:
                logger.info(f"Successfully connected to Elasticsearch at {ELASTICSEARCH_URL}.")
                # Create index if it doesn't exist
                if not ES_CLIENT.indices.exists(index=ES_INDEX):
                    logger.info(f"Elasticsearch index '{ES_INDEX}' not found. Creating...")
                    ES_CLIENT.indices.create(
                        index=ES_INDEX,
                        body={
                            "mappings": {
                                "properties": {
                                    "content": {"type": "text", "analyzer": "standard"},
                                    "file_name": {"type": "keyword"},
                                    "file_path": {"type": "keyword"},
                                    "file_size": {"type": "long"},
                                    "upload_date": {"type": "date"},
                                    "user_id": {"type": "keyword"}
                                }
                            }
                        }
                    )
                    logger.info(f"Elasticsearch index '{ES_INDEX}' created.")
                else:
                    logger.info(f"Elasticsearch index '{ES_INDEX}' already exists.")
        except ESConnectionError as e:
            logger.error(f"Failed to connect to Elasticsearch at {ELASTICSEARCH_URL}: {e}. Falling back to local search.")
            ES_CLIENT = None
        except Exception as e:
            logger.error(f"An unexpected error occurred during Elasticsearch initialization: {e}. Falling back to local search.")
            ES_CLIENT = None
    elif not ELASTICSEARCH_AVAILABLE:
        logger.info("Elasticsearch client library not installed. Using local file-based search only.")
    else: # ELASTICSEARCH_URL is not set or empty
        logger.info("Elasticsearch URL not configured. Using local file-based search only.")
        ES_CLIENT = None

def extract_text_from_document(file_path: str) -> str:
    """Extract text from a document file. Supports PDF, TXT, DOCX."""
    if not os.path.exists(file_path):
        logger.error(f"File not found for text extraction: {file_path}")
        return ""
        
    file_ext = os.path.splitext(file_path)[1].lower()
    text = ""
    try:
        if file_ext == '.pdf':
            text = extract_text_from_pdf(file_path)
        elif file_ext == '.txt':
            text = extract_text_from_txt(file_path)
        elif file_ext == '.docx':
            text = extract_text_from_docx(file_path)
        else:
            logger.warning(f"Unsupported file format for text extraction: {file_ext} for file {file_path}")
            return ""
    except Exception as e:
        logger.error(f"Error extracting text from {file_path} (format {file_ext}): {e}", exc_info=True)
        return ""
    
    if not text.strip():
        logger.warning(f"No text extracted or extracted text is empty from {file_path}")
    return text

def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from a PDF file using PyPDF2."""
    try:
        import PyPDF2
        text = ""
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            if pdf_reader.is_encrypted:
                try:
                    pdf_reader.decrypt('') # Try with empty password
                except Exception as decrypt_error:
                    logger.warning(f"Could not decrypt PDF {file_path}: {decrypt_error}. Text extraction might fail or be incomplete.")

            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n\n" # Add double newline for paragraph separation
        return text
    except ImportError:
        logger.error("PyPDF2 not installed. Cannot extract text from PDF.")
        raise # Re-raise so it can be caught by the caller
    except Exception as e:
        logger.error(f"Error extracting text from PDF {file_path}: {e}", exc_info=True)
        raise

def extract_text_from_txt(file_path: str) -> str:
    """Extract text from a TXT file, trying common encodings."""
    encodings_to_try = ['utf-8', 'latin-1', 'windows-1252']
    for encoding in encodings_to_try:
        try:
            with open(file_path, 'r', encoding=encoding) as file:
                return file.read()
        except UnicodeDecodeError:
            logger.debug(f"Failed to decode {file_path} with {encoding}.")
            continue
        except Exception as e:
            logger.error(f"Error reading TXT file {file_path} with {encoding}: {e}", exc_info=True)
            raise # Re-raise for the caller
    logger.error(f"Could not decode TXT file {file_path} with any attempted encodings.")
    return "" # Return empty if all fail

def extract_text_from_docx(file_path: str) -> str:
    """Extract text from a DOCX file using python-docx."""
    try:
        import docx # python-docx library
        doc = docx.Document(file_path)
        text = "\n".join([para.text for para in doc.paragraphs])
        return text
    except ImportError:
        logger.error("python-docx not installed. Cannot extract text from DOCX.")
        raise
    except Exception as e:
        logger.error(f"Error extracting text from DOCX {file_path}: {e}", exc_info=True)
        raise

def add_document(file_path: str, user_id: int, file_size: int = 0) -> Optional[str]:
    """
    Add a document to the QA system.
    Stores the document, extracts text, and indexes it if Elasticsearch is available.
    """
    if not UPLOAD_DIR:
        logger.error("UPLOAD_DIR not set. Cannot add document.")
        return None
    if not os.path.exists(file_path):
        logger.error(f"Source file not found: {file_path}")
        return None
    
    doc_id = str(uuid.uuid4())
    file_name = os.path.basename(file_path)
    
    try:
        text = extract_text_from_document(file_path)
        if not text.strip(): # Check if extracted text is not just whitespace
            logger.warning(f"No meaningful text extracted from document: {file_name} (ID: {doc_id}). Document not added.")
            return None
        
        upload_date = datetime.utcnow().isoformat()
        
        user_dir = os.path.join(UPLOAD_DIR, "user_documents", str(user_id))
        os.makedirs(user_dir, exist_ok=True)
        
        # Use doc_id in the stored filename to ensure uniqueness within user's folder
        # and avoid issues with users uploading files with the same name.
        sanitized_file_name = "".join(c if c.isalnum() or c in ('.', '_', '-') else '_' for c in file_name)
        dest_filename = f"{doc_id}_{sanitized_file_name}"
        dest_path = os.path.join(user_dir, dest_filename)
        
        shutil.copy2(file_path, dest_path)
        
        actual_file_size = os.path.getsize(dest_path) if os.path.exists(dest_path) else file_size

        doc_metadata = {
            "id": doc_id,
            "file_name": file_name, # Original filename for display
            "stored_file_name": dest_filename, # Actual name on disk
            "file_path": dest_path, # Full path to stored file
            "file_size": actual_file_size,
            "upload_date": upload_date,
            "user_id": str(user_id)
        }
        
        metadata_path = os.path.join(user_dir, f"{doc_id}_metadata.json")
        with open(metadata_path, 'w', encoding='utf-8') as f:
            json.dump(doc_metadata, f, indent=4)
        
        content_path = os.path.join(user_dir, f"{doc_id}_content.txt")
        with open(content_path, 'w', encoding='utf-8') as f:
            f.write(text)
        
        if ES_CLIENT:
            try:
                ES_CLIENT.index(
                    index=ES_INDEX,
                    id=doc_id, # Use the generated doc_id as ES document ID
                    document={ # 'document' kwarg for ES client v8+, 'body' for v7
                        "content": text,
                        "file_name": file_name,
                        "file_path": dest_path, # Store path for reference, not for direct ES access
                        "file_size": actual_file_size,
                        "upload_date": upload_date,
                        "user_id": str(user_id)
                    }
                )
                logger.info(f"Document '{file_name}' (ID: {doc_id}) indexed in Elasticsearch.")
            except Exception as es_err:
                logger.error(f"Failed to index document {doc_id} in Elasticsearch: {es_err}", exc_info=True)
                # Decide if this is a critical error. For now, we proceed with local storage.
        
        logger.info(f"Document added: '{file_name}' (ID: {doc_id}), stored at {dest_path}")
        return doc_id
        
    except Exception as e:
        logger.error(f"Error adding document {file_name}: {e}", exc_info=True)
        # Clean up partially created files if error occurs
        if 'dest_path' in locals() and os.path.exists(dest_path): os.remove(dest_path)
        if 'metadata_path' in locals() and os.path.exists(metadata_path): os.remove(metadata_path)
        if 'content_path' in locals() and os.path.exists(content_path): os.remove(content_path)
        return None

def query_document(doc_id: str, question: str, user_id_for_auth: Optional[int] = None) -> Dict[str, Any]:
    """
    Query a document with a question.
    Authenticates user if user_id_for_auth is provided.
    Uses Elasticsearch if available, otherwise falls back to local search.
    """
    metadata = get_document_metadata(doc_id)
    if not metadata:
        return {"error": "Document not found or metadata missing."}

    if user_id_for_auth is not None and str(metadata.get("user_id")) != str(user_id_for_auth):
        logger.warning(f"User {user_id_for_auth} attempted to query document {doc_id} owned by {metadata.get('user_id')}.")
        return {"error": "Access denied to this document."}
        
    if ES_CLIENT:
        logger.debug(f"Querying document {doc_id} via Elasticsearch for question: '{question[:50]}...'")
        es_result = search_with_elasticsearch(doc_id, question)
        if es_result and "error" not in es_result:
            return es_result
        logger.warning(f"Elasticsearch query failed or returned no results for doc {doc_id}. Falling back to local search.")

    # Fallback to local search if ES is not available or fails
    logger.debug(f"Querying document {doc_id} via local search for question: '{question[:50]}...'")
    user_doc_dir = os.path.join(UPLOAD_DIR, "user_documents", str(metadata["user_id"]))
    content_path = os.path.join(user_doc_dir, f"{doc_id}_content.txt")
    
    if not os.path.exists(content_path):
        logger.error(f"Document content file not found: {content_path}")
        return {"error": "Document content not found."}
    
    try:
        with open(content_path, 'r', encoding='utf-8') as f:
            content = f.read()
    except Exception as e:
        logger.error(f"Error reading content file {content_path}: {e}", exc_info=True)
        return {"error": "Could not read document content."}

    answer_text = simple_keyword_search(content, question)
    
    if not answer_text or answer_text == "No relevant information found.":
        return {
            "answer": "I couldn't find a specific answer to your question in this document using local search.",
            "context": "The document content doesn't seem to contain information directly related to your query based on keyword matching.",
            "source": "local_keyword_search"
        }
    
    return {
        "answer": answer_text,
        "context": get_context(content, answer_text), # Provide context around the found answer
        "source": "local_keyword_search"
    }

def search_with_elasticsearch(doc_id: str, question: str) -> Optional[Dict[str, Any]]:
    """Search for an answer within a specific document using Elasticsearch."""
    if not ES_CLIENT:
        return None
    try:
        # More sophisticated query using match for question and filter for doc_id
        query_body = {
            "query": {
                "bool": {
                    "must": [
                        {"match": {"content": question}}
                    ],
                    "filter": [
                        {"term": {"_id": doc_id}} # Filter by the specific document ID
                    ]
                }
            },
            "highlight": {
                "fields": {"content": {}}, # Highlight matches in the content field
                "pre_tags": ["<strong>"], # Optional: customize highlight tags
                "post_tags": ["</strong>"]
            },
            "size": 1 # We only need the best hit from the specified document
        }
        response = ES_CLIENT.search(index=ES_INDEX, body=query_body) # 'body' for ES v7
        
        hits = response.get("hits", {}).get("hits", [])
        if not hits:
            logger.info(f"No Elasticsearch hits for question '{question[:50]}...' in doc ID {doc_id}")
            return None
        
        best_hit = hits[0]
        answer = ""
        context = ""

        if "highlight" in best_hit and "content" in best_hit["highlight"]:
            answer = " ... ".join(best_hit["highlight"]["content"]) # Join highlighted fragments
        elif "_source" in best_hit and "content" in best_hit["_source"]:
            # Fallback if no highlight: use a snippet of the source (less ideal)
            # This part could be improved by a simple keyword search on the source if no highlight
            answer = best_hit["_source"]["content"][:500] + "..." # Basic snippet
            logger.warning(f"No highlight for ES query on doc {doc_id}, using source snippet.")
        
        if "_source" in best_hit and "content" in best_hit["_source"]:
             # Provide broader context from the source document
            full_content_for_context = best_hit["_source"]["content"]
            context = get_context(full_content_for_context, answer if answer else question, window_size=300)

        if not answer: # If answer is still empty
             logger.info(f"Elasticsearch found doc {doc_id} but no specific answer/highlight for question '{question[:50]}...'")
             return None # Or return a generic message

        return {
            "answer": answer,
            "context": context,
            "source": "elasticsearch"
        }

    except ESConnectionError as e:
        logger.error(f"Elasticsearch connection error during search: {e}", exc_info=True)
        ES_CLIENT = None # Assume connection is lost, trigger fallback for next queries
        return {"error": "Elasticsearch connection failed during search."}
    except Exception as e:
        logger.error(f"Elasticsearch search error for doc {doc_id}: {e}", exc_info=True)
        return {"error": f"Error during Elasticsearch search: {str(e)}"}


def simple_keyword_search(content: str, question: str) -> str:
    """Simple keyword-based search for an answer within the given content."""
    try:
        import re
        # NLTK is optional here, can be removed if not strictly needed or causing issues
        # from nltk.tokenize import word_tokenize
        # from nltk.corpus import stopwords
        # stop_words = set(stopwords.words('english'))
        # question_words = word_tokenize(question.lower())
        # keywords = [word for word in question_words if word.isalnum() and word not in stop_words]
        
        # Simpler keyword extraction (lowercase, alphanumeric)
        keywords = [word for word in re.findall(r'\b\w+\b', question.lower()) if len(word) > 2] # Basic keywords

        if not keywords:
            return "No keywords identified in the question for search."

        sentences = re.split(r'(?<=[.!?])\s+', content) # Split by sentence delimiters
        
        sentence_scores = []
        for i, sentence in enumerate(sentences):
            score = 0
            normalized_sentence = sentence.lower()
            for keyword in keywords:
                if keyword in normalized_sentence:
                    score += 1
            if score > 0:
                sentence_scores.append((sentence, score, i)) # Store original sentence, score, and index
        
        sentence_scores.sort(key=lambda x: x[1], reverse=True) # Sort by score
        
        if sentence_scores:
            # Return top N sentences, e.g., top 3, or a combined snippet
            top_sentences_data = sentence_scores[:3]
            top_sentences_data.sort(key=lambda x: x[2]) # Sort them by original order
            
            # Combine relevant sentences ensuring some context
            result_parts = []
            last_idx = -2
            for s_data in top_sentences_data:
                current_idx = s_data[2]
                # If sentences are not consecutive, add an ellipsis
                if current_idx > last_idx + 1 and result_parts:
                    result_parts.append("...")
                result_parts.append(s_data[0])
                last_idx = current_idx
            return " ".join(result_parts)
        
        return "No relevant information found."
    except Exception as e:
        logger.error(f"Simple keyword search error: {e}", exc_info=True)
        return "Error performing keyword search."

def get_context(content: str, answer_or_keywords: str, window_size: int = 250) -> str:
    """
    Get the context surrounding an answer or keywords within the content.
    Tries to find the first sentence of the answer/keywords.
    """
    if not content or not answer_or_keywords:
        return content[:window_size*2] + "..." if content else ""

    # Try to find the start of the answer (first few words or first sentence part)
    first_part_of_answer = answer_or_keywords.split('.')[0] # Get first sentence part
    first_part_of_answer = " ".join(first_part_of_answer.split()[:10]) # Max first 10 words

    pos = content.lower().find(first_part_of_answer.lower())
    if pos == -1: # If not found, use the beginning of the content
        pos = 0 
    
    start_context = max(0, pos - window_size)
    end_context = min(len(content), pos + len(first_part_of_answer) + window_size)
    
    # Try to align with sentence boundaries if possible (simplified)
    context_snippet = content[start_context:end_context]
    
    prefix = "..." if start_context > 0 else ""
    suffix = "..." if end_context < len(content) else ""
    
    return prefix + context_snippet + suffix

def get_document_metadata(doc_id: str) -> Optional[Dict[str, Any]]:
    """Get document metadata by searching across all user directories."""
    if not UPLOAD_DIR:
        logger.error("UPLOAD_DIR not set. Cannot get document metadata.")
        return None

    user_docs_base_dir = os.path.join(UPLOAD_DIR, "user_documents")
    if not os.path.isdir(user_docs_base_dir):
        logger.info(f"User documents base directory not found: {user_docs_base_dir}")
        return None
        
    for user_id_folder_name in os.listdir(user_docs_base_dir):
        user_dir_path = os.path.join(user_docs_base_dir, user_id_folder_name)
        if os.path.isdir(user_dir_path):
            metadata_filename = f"{doc_id}_metadata.json"
            potential_metadata_path = os.path.join(user_dir_path, metadata_filename)
            if os.path.exists(potential_metadata_path):
                try:
                    with open(potential_metadata_path, 'r', encoding='utf-8') as f:
                        metadata = json.load(f)
                        # Ensure essential keys are present
                        if metadata.get("id") == doc_id:
                            return metadata
                except json.JSONDecodeError as je:
                    logger.error(f"Error decoding JSON from metadata file {potential_metadata_path}: {je}")
                except Exception as e:
                    logger.error(f"Error reading metadata file {potential_metadata_path}: {e}", exc_info=True)
    
    logger.warning(f"Metadata for document ID {doc_id} not found across all users.")
    return None

def get_user_documents(user_id: int) -> List[Dict[str, Any]]:
    """Get all documents for a specific user."""
    if not UPLOAD_DIR:
        logger.error("UPLOAD_DIR not set. Cannot get user documents.")
        return []

    user_dir = os.path.join(UPLOAD_DIR, "user_documents", str(user_id))
    if not os.path.isdir(user_dir):
        logger.info(f"User directory not found for user {user_id} at {user_dir}")
        return []
    
    documents = []
    for filename in os.listdir(user_dir):
        if filename.endswith("_metadata.json"): # Find metadata files
            try:
                with open(os.path.join(user_dir, filename), 'r', encoding='utf-8') as f:
                    metadata = json.load(f)
                    # Format the file size for display
                    size_bytes = metadata.get("file_size", 0)
                    if size_bytes < 1024:
                        metadata["file_size_formatted"] = f"{size_bytes} B"
                    elif size_bytes < 1024 * 1024:
                        metadata["file_size_formatted"] = f"{size_bytes / 1024:.1f} KB"
                    else:
                        metadata["file_size_formatted"] = f"{size_bytes / (1024 * 1024):.1f} MB"
                    
                    # Ensure upload_date is present for sorting
                    if "upload_date" not in metadata:
                        metadata["upload_date"] = datetime.min.isoformat() # Default for sorting if missing

                    documents.append(metadata)
            except json.JSONDecodeError as je:
                logger.error(f"Error decoding JSON from {filename} for user {user_id}: {je}")
            except Exception as e:
                logger.error(f"Error reading metadata file {filename} for user {user_id}: {e}", exc_info=True)
    
    # Sort by upload date (newest first)
    documents.sort(key=lambda x: x.get("upload_date", datetime.min.isoformat()), reverse=True)
    return documents

def delete_document(doc_id: str, user_id: int) -> bool:
    """Delete a document and its associated files for a specific user."""
    if not UPLOAD_DIR:
        logger.error("UPLOAD_DIR not set. Cannot delete document.")
        return False

    metadata = get_document_metadata(doc_id) # This already searches across users
    if not metadata:
        logger.warning(f"Attempt to delete non-existent document ID: {doc_id}")
        return False
    
    # Authenticate that the document belongs to the requesting user
    if str(metadata.get("user_id")) != str(user_id):
        logger.warning(f"User {user_id} attempted to delete document {doc_id} owned by {metadata.get('user_id')}. Deletion denied.")
        return False
        
    user_specific_dir = os.path.join(UPLOAD_DIR, "user_documents", str(user_id))
    
    # Files to delete: the document itself, its metadata, and its extracted content
    # Using stored_file_name from metadata if available, otherwise construct from file_name
    stored_file_name = metadata.get("stored_file_name", f"{doc_id}_{metadata.get('file_name', 'unknown_file')}")
    paths_to_delete = [
        os.path.join(user_specific_dir, stored_file_name), # The actual document
        os.path.join(user_specific_dir, f"{doc_id}_metadata.json"),
        os.path.join(user_specific_dir, f"{doc_id}_content.txt")
    ]
    
    deleted_count = 0
    for path_to_remove in paths_to_delete:
        if os.path.exists(path_to_remove):
            try:
                os.remove(path_to_remove)
                logger.info(f"Successfully deleted file: {path_to_remove}")
                deleted_count +=1
            except OSError as e:
                logger.error(f"Error deleting file {path_to_remove}: {e}", exc_info=True)
        else:
            logger.debug(f"File not found for deletion (might be okay if optional): {path_to_remove}")

    if ES_CLIENT:
        try:
            # ignore=[404] prevents error if doc already deleted or never indexed
            ES_CLIENT.delete(index=ES_INDEX, id=doc_id, ignore=[404]) 
            logger.info(f"Document {doc_id} deleted from Elasticsearch index '{ES_INDEX}'.")
        except ESConnectionError as es_conn_err:
            logger.error(f"Elasticsearch connection error while trying to delete doc {doc_id}: {es_conn_err}. Local files may have been deleted.")
            ES_CLIENT = None # Assume connection lost
        except Exception as es_err:
            logger.error(f"Error deleting document {doc_id} from Elasticsearch: {es_err}", exc_info=True)
            # Continue, as local files might have been deleted.

    if deleted_count > 0: # If at least one file was actually deleted
        logger.info(f"Document {doc_id} (owned by user {user_id}) and its associated files deleted successfully.")
        return True
    else:
        logger.warning(f"No files were found or deleted for document ID {doc_id} for user {user_id}.")
        # If ES delete was attempted, it might have succeeded even if local files were missing.
        # Consider the operation successful if ES was available and didn't error, or if local files were removed.
        return False # Or True if ES deletion is sufficient and no local files were expected.
